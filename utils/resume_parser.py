"""
utils/resume_parser.py — Enhanced NLP-based Resume Parsing Engine
MCE Placement Prediction System
"""

import re
import json
import os
import PyPDF2
import docx
from typing import Dict, List, Optional, Tuple


# ─── Comprehensive Skill Database ─────────────────────────────────────────────
ALL_SKILLS = {
    "programming": [
        "python", "java", "c", "c++", "c#", "javascript", "typescript",
        "r", "go", "rust", "swift", "kotlin", "dart", "scala", "matlab",
        "perl", "ruby", "php", "bash", "shell", "vba", "groovy"
    ],
    "web": [
        "html", "css", "react", "angular", "vue", "node.js", "django",
        "flask", "express", "bootstrap", "tailwind", "jquery", "next.js",
        "fastapi", "spring boot", "laravel", "asp.net", "graphql", "rest api",
        "wordpress", "sass", "webpack", "vite", "svelte"
    ],
    "data": [
        "pandas", "numpy", "matplotlib", "seaborn", "plotly",
        "scikit-learn", "tensorflow", "pytorch", "keras", "xgboost",
        "power bi", "tableau", "excel", "google sheets", "looker",
        "lightgbm", "catboost", "statsmodels", "scipy", "pyspark"
    ],
    "database": [
        "sql", "mysql", "postgresql", "mongodb", "sqlite", "oracle",
        "redis", "cassandra", "firebase", "dynamodb", "elasticsearch",
        "hive", "snowflake", "bigquery", "ms sql", "mariadb", "soql"
    ],
    "cloud_devops": [
        "aws", "azure", "gcp", "docker", "kubernetes", "jenkins",
        "ansible", "terraform", "linux", "git", "github", "gitlab",
        "ci/cd", "nginx", "apache", "heroku", "vercel", "netlify",
        "bitbucket", "airflow", "kafka", "spark", "hadoop"
    ],
    "ml_ai": [
        "machine learning", "deep learning", "nlp", "computer vision",
        "neural networks", "reinforcement learning", "transformers",
        "llm", "bert", "gpt", "opencv", "yolo", "hugging face",
        "random forest", "logistic regression", "xgboost", "feature engineering",
        "model deployment", "mlops", "data augmentation", "transfer learning",
        "generative ai", "langchain", "rag", "vector database"
    ],
    "mobile": [
        "android", "ios", "react native", "flutter", "android studio",
        "xcode", "firebase", "jetpack compose", "swift ui", "kotlin"
    ],
    "tools": [
        "jupyter", "vs code", "pycharm", "postman", "figma",
        "jira", "confluence", "slack", "notion", "git",
        "linux", "windows", "macos", "vim", "intellij"
    ],
    "engineering": [
        "autocad", "solidworks", "catia", "matlab", "simulink",
        "ansys", "staad pro", "revit", "plc", "scada",
        "embedded c", "arduino", "raspberry pi", "verilog", "vhdl"
    ],
    "soft_skills": [
        "communication", "leadership", "teamwork", "problem solving",
        "time management", "critical thinking", "adaptability",
        "presentation", "project management", "agile", "scrum",
        "analytical", "research", "documentation"
    ]
}

FLAT_SKILLS = {skill for category in ALL_SKILLS.values() for skill in category}

# ─── Action Verbs ─────────────────────────────────────────────────────────────
ACTION_VERBS = [
    "developed", "built", "designed", "implemented", "created", "deployed",
    "optimized", "improved", "reduced", "increased", "managed", "led",
    "collaborated", "analyzed", "researched", "automated", "integrated",
    "trained", "evaluated", "architected", "maintained", "delivered",
    "achieved", "spearheaded", "streamlined", "engineered", "launched"
]

# ─── Section Patterns ─────────────────────────────────────────────────────────
SECTION_PATTERNS = {
    'education':      r'(education|qualification|academic|degree|university|college|school)',
    'experience':     r'(experience|employment|work history|internship|intern|job)',
    'skills':         r'(skills|technical skills|competencies|expertise|technologies)',
    'projects':       r'(projects|project work|personal projects|academic projects)',
    'certifications': r'(certifications|certificates|courses|training|achievement)',
    'contact':        r'(contact|email|phone|mobile|linkedin|github|address)',
    'summary':        r'(summary|objective|profile|about me|career objective)',
    'awards':         r'(awards|honors|achievements|recognition)',
    'publications':   r'(publications|papers|research|journal)',
    'languages':      r'(languages|spoken|written)',
}

# ─── Degree Patterns ──────────────────────────────────────────────────────────
DEGREE_PATTERNS = [
    r'b\.?tech|bachelor of technology',
    r'm\.?tech|master of technology',
    r'b\.?e\.?|bachelor of engineering',
    r'm\.?e\.?|master of engineering',
    r'bca|bachelor of computer applications',
    r'mca|master of computer applications',
    r'b\.?sc|bachelor of science',
    r'm\.?sc|master of science',
    r'bba|bachelor of business administration',
    r'mba|master of business administration',
    r'phd|doctorate|ph\.d',
    r'diploma|polytechnic',
    r'b\.?com|bachelor of commerce',
]


class ResumeParser:
    """
    Enhanced Resume Parser — In-depth extraction engine.
    Extracts: Name, Contact, Education, Experience, Projects,
    Certifications, Skills, CGPA, Links, Action Verbs, Achievements
    """

    # ── Text Extraction ────────────────────────────────────────────────────────

    def extract_text_from_pdf(self, filepath: str) -> str:
        text = ""
        try:
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
        except Exception as e:
            print(f"PDF extraction error: {e}")
        return text.strip()

    def extract_text_from_docx(self, filepath: str) -> str:
        text = ""
        try:
            doc = docx.Document(filepath)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            print(f"DOCX extraction error: {e}")
        return text.strip()

    def extract_text(self, filepath: str) -> str:
        ext = os.path.splitext(filepath)[-1].lower()
        if ext == '.pdf':
            return self.extract_text_from_pdf(filepath)
        elif ext in ['.doc', '.docx']:
            return self.extract_text_from_docx(filepath)
        return ""

    # ── Main Parse Function ────────────────────────────────────────────────────

    def parse(self, text: str) -> Dict:
        """Full in-depth parse — returns structured dict from raw resume text."""
        text_lower = text.lower()
        lines      = [l.strip() for l in text.split('\n') if l.strip()]

        experience     = self._extract_experience_detailed(text)
        projects       = self._extract_projects_detailed(text)
        certifications = self._extract_certifications(text)
        education      = self._extract_education_detailed(text)

        return {
            # Basic Info
            'name':               self._extract_name(text, lines),
            'email':              self._extract_email(text),
            'phone':              self._extract_phone(text),
            'linkedin':           self._extract_linkedin(text),
            'github':             self._extract_github(text),
            'location':           self._extract_location(text),

            # Academic
            'cgpa':               self._extract_cgpa(text),
            'tenth_percent':      self._extract_tenth(text),
            'twelfth_percent':    self._extract_twelfth(text),
            'education':          education,

            # Experience
            'experience':         experience,
            'internship_count':   self._count_internships_accurate(text),
            'internship_details': self._extract_internship_details(text),

            # Projects
            'projects':           projects,
            'project_count':      len(projects),

            # Certifications
            'certifications':     certifications,
            'certification_count': len(certifications),

            # Skills
            'skills':             self._extract_skills_advanced(text_lower, text),
            'sections_found':     self._detect_sections(text_lower),

            # Quality Indicators
            'action_verbs_found': self._count_action_verbs(text_lower),
            'has_quantified':     self._has_quantified_achievements(text),
            'word_count':         len(text.split()),
            'line_count':         len(lines),
        }

    # ── Name Extraction ────────────────────────────────────────────────────────

    def _extract_name(self, text: str, lines: List[str]) -> Optional[str]:
        """Extract name from first few lines."""
        for line in lines[:5]:
            # Skip lines with email, phone, URLs
            if re.search(r'[@\d\./]', line):
                continue
            # Skip section headers
            if re.search(r'(resume|cv|curriculum|vitae|profile)', line, re.I):
                continue
            words = line.split()
            # Name is 2-4 words, title case or all caps
            if 2 <= len(words) <= 4:
                # Check if it looks like a name
                if all(re.match(r'^[A-Za-z]+$', w) for w in words):
                    return line.title()
        return None

    # ── Contact Extraction ─────────────────────────────────────────────────────

    def _extract_email(self, text: str) -> Optional[str]:
        match = re.search(r'[\w\.\-\+]+@[\w\.\-]+\.\w{2,}', text)
        return match.group(0).lower() if match else None

    def _extract_phone(self, text: str) -> Optional[str]:
        patterns = [
            r'\+91[-\s]?[6-9]\d{9}',
            r'[6-9]\d{9}',
            r'\+\d{1,3}[-\s]?\d{10}',
        ]
        for p in patterns:
            match = re.search(p, text)
            if match:
                return match.group(0)
        return None

    def _extract_linkedin(self, text: str) -> Optional[str]:
        match = re.search(
            r'linkedin\.com/in/[\w\-]+|linkedin\.com/[\w\-/]+',
            text, re.I
        )
        return match.group(0).lower() if match else None

    def _extract_github(self, text: str) -> Optional[str]:
        match = re.search(
            r'github\.com/[\w\-]+',
            text, re.I
        )
        return match.group(0).lower() if match else None

    def _extract_location(self, text: str) -> Optional[str]:
        patterns = [
            r'\b(motihari|patna|delhi|mumbai|bangalore|bengaluru|hyderabad|'
            r'chennai|kolkata|pune|noida|gurgaon|bihar|india)\b'
        ]
        for p in patterns:
            match = re.search(p, text, re.I)
            if match:
                return match.group(0).title()
        return None

    # ── Academic Extraction ────────────────────────────────────────────────────

    def _extract_cgpa(self, text: str) -> Optional[float]:
        patterns = [
            r'cgpa[:\s]+([0-9]\.[0-9]{1,2})',
            r'gpa[:\s]+([0-9]\.[0-9]{1,2})',
            r'([0-9]\.[0-9]{1,2})\s*/\s*10\.0',
            r'([0-9]\.[0-9]{1,2})\s*/\s*10',
            r'([0-9]\.[0-9]{1,2})\s+cgpa',
            r'([0-9]\.[0-9]{1,2})\s+gpa',
        ]
        for p in patterns:
            match = re.search(p, text, re.IGNORECASE)
            if match:
                val = float(match.group(1))
                if 1.0 <= val <= 10.0:
                    return round(val, 2)
        return None

    def _extract_tenth(self, text: str) -> Optional[float]:
        patterns = [
            r'10th[^0-9]*([0-9]{2,3}\.?[0-9]*)\s*%',
            r'class\s*x[^0-9]*([0-9]{2,3}\.?[0-9]*)\s*%',
            r'matric[^0-9]*([0-9]{2,3}\.?[0-9]*)\s*%',
            r'ssc[^0-9]*([0-9]{2,3}\.?[0-9]*)\s*%',
        ]
        for p in patterns:
            match = re.search(p, text, re.IGNORECASE)
            if match:
                val = float(match.group(1))
                if 30 <= val <= 100:
                    return round(val, 1)
        return None

    def _extract_twelfth(self, text: str) -> Optional[float]:
        patterns = [
            r'12th[^0-9]*([0-9]{2,3}\.?[0-9]*)\s*%',
            r'class\s*xii[^0-9]*([0-9]{2,3}\.?[0-9]*)\s*%',
            r'hsc[^0-9]*([0-9]{2,3}\.?[0-9]*)\s*%',
            r'intermediate[^0-9]*([0-9]{2,3}\.?[0-9]*)\s*%',
            r'higher secondary[^0-9]*([0-9]{2,3}\.?[0-9]*)\s*%',
        ]
        for p in patterns:
            match = re.search(p, text, re.IGNORECASE)
            if match:
                val = float(match.group(1))
                if 30 <= val <= 100:
                    return round(val, 1)
        return None

    def _extract_education_detailed(self, text: str) -> List[Dict]:
        """Extract education with degree, institution, year, score."""
        education = []
        lines = text.split('\n')

        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            for deg_pattern in DEGREE_PATTERNS:
                if re.search(deg_pattern, line_lower):
                    edu_entry = {
                        'degree': line.strip(),
                        'institution': None,
                        'year': None,
                        'score': None,
                    }
                    # Look at next 3 lines for institution/year
                    context = ' '.join(lines[i:i+4])
                    year_m = re.search(r'(20\d{2}|19\d{2})', context)
                    if year_m:
                        edu_entry['year'] = year_m.group(0)
                    score_m = re.search(
                        r'([0-9]\.[0-9]{1,2})\s*/\s*10|([0-9]{2,3}\.?[0-9]*)\s*%',
                        context
                    )
                    if score_m:
                        edu_entry['score'] = score_m.group(0)
                    education.append(edu_entry)
                    break

        return education[:4]

    # ── Experience Extraction ──────────────────────────────────────────────────

    def _extract_experience_detailed(self, text: str) -> List[Dict]:
        """Extract experience with company, role, duration, description."""
        experiences = []
        lines = text.split('\n')
        in_section = False
        current_exp = None

        for i, line in enumerate(lines):
            line_s = line.strip()
            line_lower = line_s.lower()

            # Section start
            if re.search(
                r'^(experience|work experience|professional experience|'
                r'internship|employment)', line_lower
            ) and len(line_s) < 50:
                in_section = True
                continue

            # Section end
            if in_section and re.search(
                r'^(education|skills|projects|certif|awards|publication)',
                line_lower
            ) and len(line_s) < 50:
                if current_exp:
                    experiences.append(current_exp)
                in_section = False
                current_exp = None
                continue

            if in_section and line_s:
                # Check if this is a new experience entry
                # (has year range or company-like pattern)
                has_date = bool(re.search(
                    r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec'
                    r'|20\d{2}|present|current)', line_lower
                ))
                is_role_line = bool(re.search(
                    r'(intern|engineer|developer|analyst|manager|'
                    r'trainee|associate|consultant|researcher)',
                    line_lower
                ))

                if has_date or (is_role_line and len(line_s) < 80):
                    if current_exp:
                        experiences.append(current_exp)
                    current_exp = {
                        'role': line_s,
                        'company': None,
                        'duration': None,
                        'description': [],
                    }
                    # Extract duration
                    dur_m = re.search(
                        r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec'
                        r'|20\d{2}).{0,20}(jan|feb|mar|apr|may|jun|jul|aug|'
                        r'sep|oct|nov|dec|20\d{2}|present|current)',
                        line_lower
                    )
                    if dur_m:
                        current_exp['duration'] = dur_m.group(0)

                elif current_exp:
                    # Check if it's company name (next line after role)
                    if not current_exp['company'] and len(line_s) < 60:
                        current_exp['company'] = line_s
                    elif line_s.startswith('•') or line_s.startswith('-') or len(line_s) > 30:
                        current_exp['description'].append(line_s.lstrip('•- '))

        if current_exp:
            experiences.append(current_exp)

        return experiences[:5]

    def _extract_internship_details(self, text: str) -> List[Dict]:
        """Extract only internship entries with details."""
        all_exp = self._extract_experience_detailed(text)
        internships = []
        for exp in all_exp:
            role_lower = (exp.get('role') or '').lower()
            if re.search(
                r'intern|trainee|apprentice|summer|winter|industrial training',
                role_lower
            ):
                internships.append(exp)
        return internships

    def _count_internships_accurate(self, text: str) -> int:
        """Count internships accurately by detecting unique experience blocks."""
        internship_details = self._extract_internship_details(text)
        if internship_details:
            return min(len(internship_details), 10)

        # Fallback — count by role title lines
        lines = text.lower().split('\n')
        count = 0
        seen = set()
        for line in lines:
            line = line.strip()
            if not line or line in seen:
                continue
            if re.search(
                r'\b(summer intern|winter intern|internship trainee|'
                r'research intern|software intern|data intern|'
                r'ml intern|virtual intern|graduate intern)\b',
                line
            ):
                seen.add(line)
                count += 1
        return min(count, 10)

    # ── Projects Extraction ────────────────────────────────────────────────────

    def _extract_projects_detailed(self, text: str) -> List[Dict]:
        """Extract projects with title, tech stack, description."""
        projects = []
        lines = text.split('\n')
        in_section = False
        current_proj = None

        for line in lines:
            line_s = line.strip()
            line_lower = line_s.lower()

            # Section start
            if re.search(
                r'^(projects|project work|personal projects|academic projects|'
                r'key projects|major projects)',
                line_lower
            ) and len(line_s) < 50:
                in_section = True
                continue

            # Section end
            if in_section and re.search(
                r'^(certif|award|achievement|experience|education|skill|publication)',
                line_lower
            ) and len(line_s) < 50:
                if current_proj:
                    projects.append(current_proj)
                in_section = False
                current_proj = None
                continue

            if in_section and line_s:
                # New project — usually starts with name (not bullet)
                is_project_title = (
                    not line_s.startswith('•') and
                    not line_s.startswith('-') and
                    not line_s.startswith('*') and
                    len(line_s) < 100 and
                    not re.search(r'^(the|a|an|using|built|developed)', line_lower)
                )

                # Check for tech stack (pipe or | separator)
                has_tech = bool(re.search(r'\||\busing\b|\btech\b|\bstack\b', line_lower))

                if is_project_title and (
                    len(line_s) > 5 and not line_lower.startswith('http')
                ):
                    if current_proj and current_proj.get('title'):
                        projects.append(current_proj)
                    current_proj = {
                        'title': line_s,
                        'tech_stack': [],
                        'description': [],
                    }
                    # Extract tech from title if pipe separated
                    if '|' in line_s or '–' in line_s or '-' in line_s:
                        parts = re.split(r'\||\–|\-', line_s)
                        if len(parts) > 1:
                            current_proj['title'] = parts[0].strip()
                            tech_str = ' '.join(parts[1:])
                            current_proj['tech_stack'] = [
                                t.strip() for t in re.split(r'[,|]', tech_str)
                                if t.strip()
                            ]

                elif current_proj:
                    desc = line_s.lstrip('•-* ')
                    if desc:
                        current_proj['description'].append(desc)
                        # Extract tech from description
                        for skill in FLAT_SKILLS:
                            if skill.lower() in line_lower and skill not in current_proj['tech_stack']:
                                current_proj['tech_stack'].append(skill)

        if current_proj and current_proj.get('title'):
            projects.append(current_proj)

        return projects[:6]

    # ── Certifications ─────────────────────────────────────────────────────────

    def _extract_certifications(self, text: str) -> List[Dict]:
        """Extract certifications with name and issuer."""
        certs = []
        lines = text.split('\n')
        in_section = False

        CERT_ISSUERS = [
            'google', 'microsoft', 'aws', 'oracle', 'cisco', 'ibm',
            'coursera', 'udemy', 'nptel', 'edx', 'linkedin', 'salesforce',
            'meta', 'amazon', 'nvidia', 'tableau', 'databricks'
        ]

        for line in lines:
            line_s = line.strip()
            line_lower = line_s.lower()

            if re.search(
                r'^(certif|certificate|courses|training|achievements)',
                line_lower
            ) and len(line_s) < 50:
                in_section = True
                continue

            if in_section and re.search(
                r'^(education|skills|projects|experience|awards)',
                line_lower
            ) and len(line_s) < 50:
                in_section = False
                continue

            if in_section and line_s and len(line_s) > 5:
                cert_entry = {
                    'name': line_s.lstrip('•-* '),
                    'issuer': None
                }
                for issuer in CERT_ISSUERS:
                    if issuer in line_lower:
                        cert_entry['issuer'] = issuer.title()
                        break
                if cert_entry['name']:
                    certs.append(cert_entry)

            # Also detect inline certifications
            if re.search(
                r'certif|certified|certificate',
                line_lower
            ) and not in_section:
                for issuer in CERT_ISSUERS:
                    if issuer in line_lower:
                        cert_entry = {
                            'name': line_s.lstrip('•-* '),
                            'issuer': issuer.title()
                        }
                        if cert_entry not in certs:
                            certs.append(cert_entry)
                        break

        return certs[:10]

    # ── Skills Extraction ──────────────────────────────────────────────────────

    def _extract_skills_advanced(self, text_lower: str, text_orig: str) -> List[str]:
        """Advanced skill extraction with context awareness."""
        found = set()

        # Method 1: Exact keyword match
        for skill in FLAT_SKILLS:
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                found.add(skill)

        # Method 2: Skills section extraction
        lines = text_orig.split('\n')
        in_skills = False
        for line in lines:
            line_lower = line.lower().strip()
            if re.search(r'^(skills|technical skills|technologies|tools)', line_lower) \
                    and len(line.strip()) < 40:
                in_skills = True
                continue
            if in_skills and re.search(
                r'^(education|experience|projects|certif)', line_lower
            ) and len(line.strip()) < 40:
                in_skills = False
                continue
            if in_skills and line.strip():
                # Parse comma/pipe separated skills
                skill_items = re.split(r'[,|•:\n]', line)
                for item in skill_items:
                    item_clean = item.strip().lower()
                    if item_clean and 2 < len(item_clean) < 30:
                        # Check if it matches any known skill
                        for skill in FLAT_SKILLS:
                            if skill in item_clean or item_clean in skill:
                                found.add(skill)

        return sorted(found)

    # ── Quality Analysis ───────────────────────────────────────────────────────

    def _count_action_verbs(self, text_lower: str) -> int:
        """Count action verbs used in resume."""
        count = 0
        for verb in ACTION_VERBS:
            if re.search(r'\b' + verb + r'\b', text_lower):
                count += 1
        return count

    def _has_quantified_achievements(self, text: str) -> bool:
        """Check if resume has quantified achievements (numbers + %)."""
        patterns = [
            r'\d+\s*%',                    # percentage
            r'\d+\s*(x|times|fold)',       # multiplier
            r'\$\s*\d+',                   # dollar amount
            r'\d+\s*(users|students|records|requests|hours|days|months)',
            r'(reduced|improved|increased|optimized|saved).{0,30}\d+',
        ]
        for p in patterns:
            if re.search(p, text, re.I):
                return True
        return False

    # ── Section Detection ──────────────────────────────────────────────────────

    def _detect_sections(self, text_lower: str) -> List[str]:
        found = []
        for section, pattern in SECTION_PATTERNS.items():
            if re.search(pattern, text_lower):
                found.append(section)
        return found


# ─── Singleton ─────────────────────────────────────────────────────────────────
resume_parser = ResumeParser()
